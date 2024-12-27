from abc import ABC, abstractmethod
from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.textinput import TextInput
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.popup import Popup
from docx import Document
from openpyxl import Workbook
import os
class Delivery(ABC):
    def __init__(self, distance, time, weight, volume=0, delivery_type="standard"):
        self._distance = distance
        self._time = time
        self._weight = weight
        self._volume = volume
        self.delivery_type = delivery_type

    @property
    def distance(self):
        print(self._distance)
        return self._distance

    @distance.setter
    def setdistance(self,  value):
        try:
            float(value)
        except ValueError:
            print("invalid value")
        self._distance = float(value)
        

    @property
    def weight(self):
        return self._distance
    @weight.setter
    def setweight(self,value):
        try:
            float(value)
        except ValueError:
            print("invalid value")
        self._weight = float(value)
    
    @property
    def time(self):
        return self._distance
    @time.setter
    def setweight(self,value):
        try:
            float(value)
        except ValueError:
            print("invalid value")
        self._time = float(value)

    @property
    def volume(self):
        return self._distance
    @volume.setter
    def setvolume(self,value):
        try:
            float(value)
        except ValueError:
            print("invalid value")
        self._volume = float(value)

    @abstractmethod
    def calculate_price(self):
        pass

    def __str__(self):
        return f"{self.__class__.__name__}: {vars(self)}"

class Letter(Delivery):
    def calculate_price(self):
        return self.distance*0.2 + self.weight + self.time*0.1

class Parcel(Delivery):
    def calculate_price(self):
        return self.distance*0.3 + self.weight*0.1 + self.time*0.1 + self.volume*0.1

class Package(Delivery):
    def calculate_price(self):
        q = 1.0
        if self.delivery_type == "express":
            q = 1.5
        return self.distance*0.2 + self.weight + self.time*0.1 + self.volume*0.1

class DeliveryApp(App):
    def build(self):
        self.last_result = {}

        self.main_layout = BoxLayout(orientation="vertical")

        self.main_layout.add_widget(Label(text="what do u want to deliver:"))
        self.delivery_type_spinner = Spinner(text="letter", values=["letter", "parcel", "package"])
        self.main_layout.add_widget(self.delivery_type_spinner)

        self.inputs = {}
        for label_text in ["distance", "time", "weight", "volume", "Delivery Type (express?)"]:
            self.main_layout.add_widget(Label(text=label_text))
            input_field = TextInput(multiline=False)
            self.main_layout.add_widget(input_field)
            self.inputs[label_text.lower()] = input_field

        calculate_button = Button(text="calculate")
        calculate_button.bind(on_press=self.calculate)
        self.main_layout.add_widget(calculate_button)

        save_button = Button(text="save report")
        save_button.bind(on_press=self.save_report)
        self.main_layout.add_widget(save_button)

        self.result_label = Label(text="price: ")
        self.main_layout.add_widget(self.result_label)

        return self.main_layout
    

    def calculate(self, instance):
        distance = self.inputs["distance"].text
        time = self.inputs["time"].text
        weight = self.inputs["weight"].text
        volume = self.inputs["volume"].text if self.inputs["volume"].text else 0
        delivery_type = self.inputs["delivery type (express?)"].text.lower()

        delivery_class = {"letter": Letter, "parcel": Parcel, "package": Package}[self.delivery_type_spinner.text]
        delivery = delivery_class(distance,time,weight,delivery_type=delivery_type)
        delivery.distance = distance
        delivery.time = time
        delivery.weight = weight
        delivery.volume = volume
        price = delivery.calculate_price()

        self.last_result = {"type": self.delivery_type_spinner.text, "price": price}
        self.result_label.text = f"Price: {price:.2f}"
        # f = Document()
        # f.save("test.docx")

    def save_report(self, instance):

        filetypes = [("Word Document", "*.docx"), ("Excel Workbook", "*.xlsx")]
        f = open("report.docx")
        if not self.last_result:
            popup = Popup(title="save report", content=Label(text="no results to save"), size_hint=(0.5, 0.5))
            popup.open()
        else:
            
            doc = Document()
            for key, value in self.last_result.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.save("report.docx")
            popup = Popup(title="save report", content=Label(text="file saved"), size_hint=(0.5, 0.5))
            popup.open()
        return

    def show_error(self, message):
        popup = Popup(title="error", content=Label(text=message), size_hint=(0.5, 0.5))
        popup.open()


DeliveryApp().run()